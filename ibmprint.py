from PIL import Image
import emoji
from docx import Document
from docx.shared import Pt, Inches
from html.parser import HTMLParser
import os
from uuid import uuid4
from tinydb import where
import time
import threading
import segno

def suuid():
    return str(uuid4())

def GetKey(d, val):
    for key, value in d.items():
        if val == value:
            return key
    return "Key Not Found"

class pronter:
    def __init__(self, print_directory, queue_directory, fonts, tags, db, max_time_pending):
        self.pd = print_directory
        self.qd = queue_directory
        self.fonts = fonts
        self.tags = tags
        self.db = db
        self.max_time_pending = float(max_time_pending)

    def flushqueue(self):
        #TODO: There must be a better way to iterate of the items

        #Remove DB entries that have already been removed or executed
        #TODO: Also do the opposite; remove files unlisted in DB
        queue = self.db.table("printqueue").all()
        for job in queue:
            if not os.path.isfile(self.qd + job["file"]):
                self.db.table("printqueue").remove(where("uuid") == job["uuid"])

        #Compile list of print jobs
        toprint = []
        for job in queue:
            if self.db.table("users").search(where("id") == job["id"]): #check if the job's sender is allowed
                toprint.append(job)
            else:
                if time.time() - job["time"] > self.max_time_pending:
                    os.remove(self.qd + job["file"])

        def do_flush(self, toprint):
            for job in toprint:
                while os.listdir(self.pd):
                    time.sleep(.25) #TODO: consider some escape clause for if trash piles up in the folder
                try:
                    os.rename(self.qd + job["file"], self.pd + job["file"])
                except FileNotFoundError:
                    pass
        
        queue = self.db.table("printqueue").all()
        threading.Thread(target=do_flush, args=(self,toprint)).start()
        
    def new_job(self, thing, doctype, caption=None, user_id=None):
        #TODO: make better exception handling
        try:
            uuid, filename = self.prepare_document(thing, doctype, caption)
            self.db.table("printqueue").insert({"uuid":uuid, "file":filename, "id":str(user_id), "time":time.time()})
            return False,
        except Exception as e:
            return True, str(e)

    def prepare_document(self, thing, doctype, caption):
        uuid = suuid()
        
        if doctype == 'text':
            document = Document()
            parseToParagraph(self.tags, self.fonts, document.add_paragraph()).feed(thing)
            filename = uuid + ".docx"
            document.save(self.qd + filename)

        elif doctype == 'photo' or doctype == 'png':
            document = Document()
            imgf = self.save_image(thing)
            document.add_picture(imgf, width=Inches(2.75))
            os.remove(imgf)
            if caption is not None:
                parseToParagraph(self.tags, self.fonts, document.add_paragraph()).feed(caption)
            filename = uuid + ".docx"
            document.save(self.qd + filename)
        
        elif doctype == 'doc':
            filename = uuid + "_direct" + ".doc"
            thing.download(self.qd + filename)

        elif doctype == 'docx':
            filename = uuid + "_direct" + ".docx"
            thing.download(self.qd + filename)
            
        elif doctype == 'qr':
            document = Document()
            imgf = self.qd + suuid() + ".png"
            segno.make(thing, micro=False).save(imgf, scale=50)
            document.add_picture(imgf, width=Inches(2.75))
            os.remove(imgf)
            filename = uuid + ".docx"
            document.save(self.qd + filename)

        return uuid, filename
    
    def save_image(self, image):
        uuid = suuid()
        fullname = self.qd + uuid + ".temp"
        image.download(fullname)
        image_file = Image.open(fullname)

        #transparency to white
        if 'A' in image_file.getbands():
            image_file = image_file.convert('RGBA')
            datas = image_file.getdata()
            newData = []
            for item in datas:
                if item[3] == 0:
                    newData.append((255, 255, 255, 0))
                elif item[3] != 255:
                    R = item[3]/255
                    bc = 255 - item[3]
                    newData.append((int(R*item[0] + bc), int(R*item[1] + bc), int(R*item[2] + bc), 0))
                else:
                    newData.append(item)
            image_file.putdata(newData)
        
        #to black and white
        image_file = image_file.convert('1') 
        
        finalname = fullname
        finalname = fullname.replace('temp','bmp')
        image_file.save(finalname)
        os.remove(fullname)

        return finalname


class parseToParagraph(HTMLParser):
    def __init__(self, tags, fonts, paragraph):
        super().__init__()
        self.reset()
        self.tags = tags
        self.fonts = fonts
        self.paragraph = paragraph

    def handle_starttag(self, tag, attrs):
        try:
            sleutel = GetKey(self.tags, tag.upper())
            self.font = self.fonts[sleutel]
        except KeyError:
            self.font = self.fonts["normal"]
    
    def handle_endtag(self, tag):
        self.font = self.fonts["normal"]

    def handle_data(self, text):
        try:
            if self.font is None:
                self.font = self.fonts["normal"]
        except AttributeError:
            self.font = self.fonts["normal"]
        
        emojis = emoji.emoji_lis(text)
        if len(emojis) != 0:
            Nemo = len(emojis)
            Ntxt = len(text)
            #divide the text up into segments enclosed by emoji
            segments = []
            for i in range(Nemo):
                if emojis[i]["location"] == 0:
                    #add emojus
                    segments.append((0, 1, True))
                    #add text following emojus
                    first = 1
                    try:
                        last = emojis[i+1]["location"]
                    except IndexError:
                        last = Ntxt
                    segments.append((first, last, False))
                elif emojis[i]["location"] == Ntxt-1:
                    #add text preceding first emojus
                    if i == 0:
                        segments.append((0, emojis[i]["location"], False))
                    #add emojus
                    segments.append((Ntxt-1, Ntxt, True))
                else:
                    #add text preceding first emojus
                    if i == 0:
                        segments.append((0, emojis[i]["location"], False))
                    #add emojus
                    segments.append((emojis[i]["location"], emojis[i]["location"]+1, True))
                    #add text following emojus
                    first = emojis[i]["location"]+1
                    try:
                        last = emojis[i+1]["location"]
                    except IndexError:
                        last = Ntxt
                    segments.append((first, last, False))
            #add all segments to the document
            for seg in segments:
                if seg[2]:
                    run = self.paragraph.add_run(text[seg[0]:seg[1]])
                    run.font.size = Pt(10)
                    run.font.name = self.fonts["emoji"]
                else:
                    run = self.paragraph.add_run(text[seg[0]:seg[1]])
                    run.font.name = self.font
                    run.font.size = Pt(10)
                    run.bold = True   
        else:
            run = self.paragraph.add_run(text)
            run.font.name = self.font
            run.font.size = Pt(10)
            run.bold = True
