"""
Usage:

curl -X POST "http://localhost:8000/chat" \
     -H "Authorization: strikt-geheim" \
     -H "Content-Type: multipart/form-data" \
     -F "file=@./some-image.png" \
     -F "text=funny text goes here"
"""

from fastapi import FastAPI, HTTPException, Header, File, Form, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from PIL import Image
import io
from docx import Document
from docx.shared import Pt, Inches
import validators
import segno
import socket
import argparse
import uvicorn
import emoji
from html.parser import HTMLParser
import os


IMG_WIDTH = Inches(2.75)
IMG_FORMAT = "png"


app = FastAPI()


app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


def process_image(image_file: Image) -> Image:
    # transparency to white
    if "A" in image_file.getbands():
        image_file = image_file.convert("RGBA")
        datas = image_file.getdata()
        newData = []
        for item in datas:
            if item[3] == 0:
                newData.append((255, 255, 255, 0))
            elif item[3] != 255:
                R = item[3] / 255
                bc = 255 - item[3]
                newData.append(
                    (
                        int(R * item[0] + bc),
                        int(R * item[1] + bc),
                        int(R * item[2] + bc),
                        0,
                    )
                )
            else:
                newData.append(item)
        image_file.putdata(newData)

    # to black and white
    return image_file.convert("1")


class parseToParagraph(HTMLParser):
    TAGS = {
        "fast": "F",
        "barcode": "BAR",
    }

    FONTS = {
        "normal": "Courier New",
        "emoji": "Segoe UI Emoji",
        "fast": "Font A",
        "barcode": "JAN 13 (EAN-13)",
    }

    SIZE = 10

    @staticmethod
    def _key_from_value(d, val):
        for key, value in d.items():
            if val == value:
                return key
        return "Key Not Found"

    def __init__(self, paragraph):
        super().__init__()
        self.reset()
        self.paragraph = paragraph
        self.font = self.FONTS["normal"]

    def handle_starttag(self, tag, attrs):
        try:
            sleutel = self._key_from_value(self.TAGS, tag.upper())
            self.font = self.FONTS[sleutel]
        except KeyError:
            self.font = self.FONTS["normal"]

    def handle_endtag(self, tag):
        self.font = self.FONTS["normal"]

    def handle_data(self, text):
        emojis = emoji.emoji_list(text)
        if len(emojis) != 0:
            Nemo = len(emojis)
            Ntxt = len(text)
            # divide the text up into segments enclosed by emoji
            segments = []
            for i in range(Nemo):
                if emojis[i]["match_start"] == 0:
                    # add emojus
                    segments.append((0, 1, True))
                    # add text following emojus
                    first = 1
                    try:
                        last = emojis[i + 1]["match_end"]
                    except IndexError:
                        last = Ntxt
                    segments.append((first, last, False))
                elif emojis[i]["match_end"] == Ntxt - 1:
                    # add text preceding first emojus
                    if i == 0:
                        segments.append((0, emojis[i]["match_end"], False))
                    # add emojus
                    segments.append((Ntxt - 1, Ntxt, True))
                else:
                    # add text preceding first emojus
                    if i == 0:
                        segments.append((0, emojis[i]["match_end"], False))
                    # add emojus
                    segments.append(
                        (emojis[i]["match_start"], emojis[i]["match_end"], True)
                    )
                    # add text following emojus
                    first = emojis[i]["match_end"]
                    try:
                        last = emojis[i + 1]["match_end"]
                    except IndexError:
                        last = Ntxt
                    segments.append((first, last, False))
            # add all segments to the document
            for first, last, is_emojus in segments:
                run = self.paragraph.add_run(text[first:last])
                run.font.size = Pt(self.SIZE)
                if is_emojus:
                    run.font.name = self.FONTS["emoji"]
                else:
                    run.font.name = self.font
                    run.bold = True
        else:
            run = self.paragraph.add_run(text)
            run.font.name = self.font
            run.font.size = Pt(self.SIZE)
            run.bold = True


def authorized(authorization: str) -> bool:
    return authorization == os.environ["SUREMARK_SECRET"]


def print_document(document):
    bio = io.BytesIO()
    document.save(bio)
    with socket.create_connection((args.nchost, args.ncport)) as sock:
        sock.sendall(bio.getvalue())


@app.post("/auth")
def auth(authorization: str = Header(None)):
    if not authorized(authorization):
        raise HTTPException(status_code=403, detail="Unauthorized")
    return {"detail": "lekker man"}


@app.post("/chat")
def handle_message(
    authorization: str = Header(None),
    text: str = Form(""),
    file: UploadFile = File(None),
):
    if not authorized(authorization):
        raise HTTPException(status_code=403, detail="Unauthorized")
    
    if not text and not file:
        raise HTTPException(status_code=400, detail="Nothing to print")

    document = Document()

    # Text-only form content is url, print QR code
    if (
        text
        and not file
        and validators.url(text if "://" in text else f"https://{text}")
    ):
        img_byte_arr = io.BytesIO()
        segno.make(text, micro=False).save(img_byte_arr, kind=IMG_FORMAT, scale=50)
        document.add_picture(img_byte_arr, width=IMG_WIDTH)
    
    # Print normal document, possibly with image
    else:
        if file:
            img_bytes = io.BytesIO(file.file.read())
            try:
                Image.open(img_bytes).verify()
            except Exception:
                raise HTTPException(
                    status_code=415, detail="Uploaded file does not look like an image."
                )
            else:
                img = Image.open(img_bytes)
                img = process_image(img)
                img_byte_arr = io.BytesIO()
                img.save(img_byte_arr, format=IMG_FORMAT)
                document.add_picture(img_byte_arr, width=IMG_WIDTH)
        if text:
            parseToParagraph(document.add_paragraph()).feed(text)

    print_document(document)
    return {"detail": "ok"}


if __name__ == "__main__":
    parser = argparse.ArgumentParser(
        description="Receive multi-part forms with text and images from authorized users, generate a docx-document and cat that to some host."
    )
    parser.add_argument("nchost", help="Host to send the documents to")
    parser.add_argument("--ncport", type=int, help="Port of target host", default=8000)
    parser.add_argument("--listen", help="Listen address", default="127.0.0.1")
    parser.add_argument("--port", type=int, help="Port to host on", default=8000)

    args = parser.parse_args()
    uvicorn.run(app, host=args.listen, port=args.port)
